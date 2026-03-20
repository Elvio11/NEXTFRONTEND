import sys
import json
import base64
from io import BytesIO
from mcp.server.fastmcp import FastMCP
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

# Initialize FastMCP server
mcp = FastMCP("DocxGenerator")

def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    # Horizontal rule via paragraph border — simple underline approach
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)


def _build_docx(tailored: dict) -> Document:
    """Convert the tailored resume dict into a formatted .docx Document."""
    doc = Document()

    # Tight margins
    section = doc.sections[0]
    section.top_margin    = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin   = Pt(54)
    section.right_margin  = Pt(54)

    # Name
    name_para = doc.add_paragraph()
    name_run  = name_para.add_run(tailored.get("name", ""))
    name_run.bold      = True
    name_run.font.size = Pt(16)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact
    contact = f"{tailored.get('email', '')}  |  {tailored.get('phone', '')}"
    ct = doc.add_paragraph(contact)
    ct.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Summary
    if tailored.get("summary"):
        _add_section_heading(doc, "PROFESSIONAL SUMMARY")
        doc.add_paragraph(tailored["summary"])

    # Skills
    if tailored.get("skills"):
        _add_section_heading(doc, "SKILLS")
        skills_text = "  •  ".join(tailored["skills"])
        doc.add_paragraph(skills_text)

    # Experience
    if tailored.get("experience"):
        _add_section_heading(doc, "EXPERIENCE")
        for exp in tailored["experience"]:
            title_para = doc.add_paragraph()
            title_run  = title_para.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
            title_run.bold = True
            doc.add_paragraph(exp.get("dates", ""))
            for bullet in exp.get("bullets", []):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(bullet)

    # Education
    if tailored.get("education"):
        _add_section_heading(doc, "EDUCATION")
        for edu in tailored["education"]:
            edu_para = doc.add_paragraph()
            edu_para.add_run(
                f"{edu.get('degree', '')} — {edu.get('institution', '')} ({edu.get('year', '')})"
            )

    # Certifications
    if tailored.get("certifications"):
        _add_section_heading(doc, "CERTIFICATIONS")
        for cert in tailored["certifications"]:
            doc.add_paragraph(f"• {cert}")

    return doc


@mcp.tool()
def generate_docx(payload_json: str) -> str:
    """
    Generate a tailored DOCX resume from JSON data.
    Takes a JSON string representing the tailored resume structure.
    Returns: Base64 encoded DOCX file content.
    """
    try:
        tailored = json.loads(payload_json)
        doc = _build_docx(tailored)
        
        # Save to memory boundary
        stream = BytesIO()
        doc.save(stream)
        
        # Convert to base64 for reliable MCP transport
        b64_content = base64.b64encode(stream.getvalue()).decode('utf-8')
        
        return json.dumps({
            "status": "success", 
            "content_base64": b64_content
        })
    except Exception as exc:
        return json.dumps({
            "status": "error",
            "message": str(exc)
        })

def _build_pdf(tailored: dict) -> bytes:
    """Convert tailored resume dict into a PDF byte string."""
    pdf = FPDF(unit="pt", format="A4")
    pdf.add_page()
    pdf.set_margins(left=54, top=36, right=54)
    pdf.set_auto_page_break(auto=True, margin=36)

    # Add standard font
    pdf.set_font("Helvetica", "B", 16)
    
    # Name
    name = tailored.get("name", "")
    pdf.cell(0, 20, name, align="C", new_x="LMARGIN", new_y="NEXT")
    
    # Contact
    pdf.set_font("Helvetica", "", 10)
    contact = f"{tailored.get('email', '')}  |  {tailored.get('phone', '')}"
    pdf.cell(0, 15, contact, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)

    def _add_section_heading(text: str):
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(26, 86, 219) # RGBColor(0x1A, 0x56, 0xDB)
        pdf.cell(0, 15, text, new_x="LMARGIN", new_y="NEXT")
        # simple underline
        pdf.set_draw_color(26, 86, 219)
        pdf.line(pdf.get_x(), pdf.get_y(), pdf.w - 54, pdf.get_y())
        pdf.ln(4)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 10)

    # Summary
    if tailored.get("summary"):
        _add_section_heading("PROFESSIONAL SUMMARY")
        pdf.multi_cell(0, 12, tailored["summary"], new_x="LMARGIN", new_y="NEXT")
        pdf.ln(8)

    # Skills
    if tailored.get("skills"):
        _add_section_heading("SKILLS")
        skills_text = "  #  ".join(tailored["skills"])
        pdf.multi_cell(0, 12, skills_text, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(8)

    # Experience
    if tailored.get("experience"):
        _add_section_heading("EXPERIENCE")
        for exp in tailored["experience"]:
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 12, f"{exp.get('title', '')} - {exp.get('company', '')}", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "I", 10)
            pdf.cell(0, 12, exp.get("dates", ""), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            for bullet in exp.get("bullets", []):
                # Pseudo bullet point
                pdf.multi_cell(0, 12, f"\x95 {bullet}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(5)

    # Education
    if tailored.get("education"):
        _add_section_heading("EDUCATION")
        for edu in tailored["education"]:
            pdf.multi_cell(0, 12, f"{edu.get('degree', '')} - {edu.get('institution', '')} ({edu.get('year', '')})", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)

    # Certifications
    if tailored.get("certifications"):
        _add_section_heading("CERTIFICATIONS")
        for cert in tailored["certifications"]:
            pdf.cell(0, 12, f"\x95 {cert}", new_x="LMARGIN", new_y="NEXT")
            
    return pdf.output()


@mcp.tool()
def generate_pdf(payload_json: str) -> str:
    """
    Generate a tailored PDF resume from JSON data.
    Takes a JSON string representing the tailored resume structure.
    Returns: Base64 encoded PDF file content.
    """
    try:
        tailored = json.loads(payload_json)
        pdf_bytes = _build_pdf(tailored)
        
        b64_content = base64.b64encode(pdf_bytes).decode('utf-8')
        
        return json.dumps({
            "status": "success", 
            "content_base64": b64_content
        })
    except Exception as exc:
        return json.dumps({
            "status": "error",
            "message": str(exc)
        })


if __name__ == "__main__":
    mcp.run()
