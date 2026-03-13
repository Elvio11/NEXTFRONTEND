"""
skills/resume_tailor.py
Tailors a user's parsed resume to a specific job description using Sarvam-M Think.

Reads the parsed resume from FluxShare (/storage/parsed-resumes/{user_id}.json.gz).
Calls Sarvam-M Think to reorder/reframe bullets using JD keyword alignment.
Generates a python-docx .docx file written to /storage/tailored-resumes/{user_id}/{job_id}.docx.

THE ABSOLUTE CONSTRAINT: NEVER fabricate skills, experience, certifications, or achievements.
Only reorder and reframe what already exists in the resume. Missing skills stay missing.
"""

import json
from io import BytesIO
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from llm.sarvam import sarvam, SarvamUnavailableError
from skills.humanizer_prompt import HUMANIZER_GUIDELINES
from skills.storage_client import get_json_gz, put_bytes


# ─── Tailoring Prompt ──────────────────────────────────────────────────────────

def _build_tailor_prompt(resume: dict, job: dict) -> str:
    return f"""You are an expert resume tailor. Rewrite the resume to maximise fit for this specific job.

ABSOLUTE RULES (never violate):
- NEVER add skills, certifications, or experience that are not in the original resume
- NEVER change dates, company names, job titles, or education institutions
- ONLY reorder bullets and reframe using JD terminology where synonyms exist
- Missing skills stay missing — do not imply them

JOB DETAILS:
Title:   {job.get('title', '')}
Company: {job.get('company', '')}
Required skills: {', '.join(job.get('required_skills', [])[:20])}
JD summary: {job.get('jd_summary', job.get('raw_jd', ''))[:1000]}

ORIGINAL RESUME:
{json.dumps(resume, indent=2)[:3000]}

{HUMANIZER_GUIDELINES}

OUTPUT FORMAT — Return ONLY a JSON object with this exact structure:
{{
  "name":         "<full name from resume>",
  "email":        "<email>",
  "phone":        "<phone>",
  "summary":      "<rewritten 2-3 sentence summary aligned to this role>",
  "skills":       ["<required skills first, then others — no new skills>"],
  "experience": [
    {{
      "title":   "<exact job title — do not change>",
      "company": "<exact company — do not change>",
      "dates":   "<exact dates — do not change>",
      "bullets": ["<reordered/reframed bullet — most relevant to THIS job first>"]
    }}
  ],
  "education": [
    {{
      "degree":      "<exact degree>",
      "institution": "<exact institution>",
      "year":        "<year>"
    }}
  ],
  "certifications": ["<certifications from original only>"]
}}"""


# ─── DOCX Generator ────────────────────────────────────────────────────────────

def _build_docx(tailored: dict) -> Document:
    """Convert the tailored resume dict into a formatted .docx Document."""
    doc = Document()

    # Tight margins
    section = doc.sections[0]
    section.top_margin    = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin   = Pt(54)
    section.right_margin  = Pt(54)

    # Name
    name_para = doc.add_paragraph()
    name_run  = name_para.add_run(tailored.get("name", ""))
    name_run.bold      = True
    name_run.font.size = Pt(16)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact
    contact = f"{tailored.get('email', '')}  |  {tailored.get('phone', '')}"
    ct = doc.add_paragraph(contact)
    ct.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Summary
    if tailored.get("summary"):
        _add_section_heading(doc, "PROFESSIONAL SUMMARY")
        doc.add_paragraph(tailored["summary"])

    # Skills
    if tailored.get("skills"):
        _add_section_heading(doc, "SKILLS")
        skills_text = "  •  ".join(tailored["skills"])
        doc.add_paragraph(skills_text)

    # Experience
    if tailored.get("experience"):
        _add_section_heading(doc, "EXPERIENCE")
        for exp in tailored["experience"]:
            title_para = doc.add_paragraph()
            title_run  = title_para.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
            title_run.bold = True
            doc.add_paragraph(exp.get("dates", ""))
            for bullet in exp.get("bullets", []):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(bullet)

    # Education
    if tailored.get("education"):
        _add_section_heading(doc, "EDUCATION")
        for edu in tailored["education"]:
            edu_para = doc.add_paragraph()
            edu_para.add_run(
                f"{edu.get('degree', '')} — {edu.get('institution', '')} ({edu.get('year', '')})"
            )

    # Certifications
    if tailored.get("certifications"):
        _add_section_heading(doc, "CERTIFICATIONS")
        for cert in tailored["certifications"]:
            doc.add_paragraph(f"• {cert}")

    return doc


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    # Horizontal rule via paragraph border — simple underline approach
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)


# ─── Main Entry ────────────────────────────────────────────────────────────────

async def tailor_resume(user_id: str, job: dict) -> str:
    """
    Tailor the user's resume to the given job and write the .docx to FluxShare.

    Args:
        user_id: User UUID
        job: dict with keys: id, title, company, required_skills, jd_summary, raw_jd

    Returns:
        Storage path string: /storage/tailored-resumes/{user_id}/{job_id}.docx

    Raises:
        SarvamUnavailableError: if Sarvam-M is down — agent catches and returns skipped
        FileNotFoundError: if parsed resume doesn't exist on FluxShare
    """
    job_id = str(job.get("id", "unknown"))

    try:
        resume = await get_json_gz(f"parsed-resumes/{user_id}.json.gz")
    except Exception as exc:
        raise FileNotFoundError(f"Parsed resume not found: {exc}")

    prompt  = _build_tailor_prompt(resume, job)
    raw     = await sarvam.complete(prompt, mode="think")

    # Parse Sarvam-M JSON response
    try:
        # Strip markdown fences if present
        clean = raw.strip()
        if clean.startswith("```"):
            clean = clean.split("```")[1]
            if clean.startswith("json"):
                clean = clean[4:]
        tailored = json.loads(clean.strip())
    except (json.JSONDecodeError, IndexError) as exc:
        raise SarvamUnavailableError(f"Sarvam returned non-JSON tailor response: {exc}") from exc

    # Build and write DOCX
    output_path = f"tailored-resumes/{user_id}/{job_id}.docx"

    doc = _build_docx(tailored)
    stream = BytesIO()
    doc.save(stream)
    await put_bytes(output_path, stream.getvalue())

    return output_path
